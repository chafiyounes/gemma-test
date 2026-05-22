"""Category-aware document loader + BM25 retrieval.

Layout (per category folder under ``data/documents/<name>/``):

1. Try **Markdown** ``data/documents_md/<name>/**/*.md`` — used **only if** at least one
   file indexes (non-empty, tokenizable). Otherwise fall through (so empty/stale MD
   mirrors do not hide a real ``data/documents/<name>/`` corpus).

2. Else try **plain text** ``data/documents_txt/<name>/**/*.txt`` with the same rule.

3. Else **Word/PDF** under ``data/documents/<name>/`` (``*.docx`` and optional ``pdf/*.pdf``).

Main-body only: headers, footers, and images are not extracted. Files placed
directly in data/documents/ (without a subfolder) are ignored.
"""
from __future__ import annotations

import logging
import math
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from app_config.settings import settings
from core.logigrammes_store import append_to_document_text

REPO_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = REPO_ROOT / "data" / "documents"
DOCS_MD_DIR = REPO_ROOT / "data" / "documents_md"
DOCS_TXT_DIR = REPO_ROOT / "data" / "documents_txt"

from core.sop_text_clean import clean_sop_markdown, collapse_whitespace  # noqa: E402
from core.docx_to_md import convert_docx_to_markdown  # noqa: E402

logger = logging.getLogger(__name__)

_TOKEN_RE = re.compile(r"[A-Za-zÀ-ÿ0-9\u0600-\u06FF]+", re.UNICODE)

# If hints didn't fire, Darija / mixed messages often have no French lemmas for BM25.
_DARIJA_LEX_BOOST = re.compile(
    r"\b("
    r"bghay|bgha|bghit|bghiti|chno|chnowa|ntebi3|ntebi3ohom|dyal|dial|diali|"
    r"walakin|wakha|katlab|ybdel|ybeddel|flivraison|f\s+livraison"
    r")\b",
    re.IGNORECASE,
)

def _tokenize(text: str) -> List[str]:
    return [t.lower() for t in _TOKEN_RE.findall(text)]


# BM25 is lexical: French/Darija questions may miss French SOP terms unless we
# append French synonyms. **English questions must stay English-only** (no
# French keyword stuffing) — rely on normal English terms + full-category inject.
_LOGISTICS_EN_FR: tuple[tuple[str, str], ...] = (
    ("delivery", "livraison livrer livraisons"),
    ("deliver", "livrer livraison"),
    ("region", "région destination villes zone couverture"),
    ("vendor", "vendeur vendeurs"),
    ("customer", "client clients client final"),
    ("refund", "remboursement"),
    ("damaged", "endommagé dommage incident transport"),
    ("pickup", "ramassage collecte pickup"),
    ("verification", "vérification vérifier"),
    ("verify", "vérifier"),
    ("system", "système plateforme"),
    ("does not appear", "n'apparaît pas n'apparaissent pas"),
    ("specific", "spécifique"),
    ("asking if", "question demande"),
    ("final client", "client final"),
    ("phone", "téléphone téléphone coordonnées client contact"),
    ("mobile", "téléphone portable contact client"),
    ("number", "numéro téléphone contact client coordonnées"),
    ("change", "modification changement coordonnées procédure"),
    ("update", "mise à jour modification changement coordonnées"),
    ("stock", "stock entrepôt entrée sortie mouvement inventaire produit article demande catalogue"),
    ("restock", "entrée stock réapprovisionnement produit mouvement"),
    ("inventory", "stock inventaire entrepôt produit article"),
    ("product", "produit produits article stock catalogue ajouter"),
    ("products", "produits stock articles catalogue ajouter"),
    ("warehouse", "entrepôt stock magasinage"),
    ("request", "demande procédure support client entrée sortie activation mouvement"),
)


# Extra BM25 terms when the query is already in French (or Darija + French lexique).
_LOGISTICS_FR_HINTS: tuple[tuple[str, str], ...] = (
    ("distribué", "statut colis livré réception livraison"),
    ("injoignable", "tentative livraison livreur contact inaccessible"),
    ("retour en cours", "retour colis entrepôt expéditeur statut"),
    ("retour à l'entrepôt", "retour entrepôt colis statut"),
    ("en attente de retour", "retour colis retour expéditeur"),
    ("facturé", "facture facturation statut"),
    ("facture", "facturation statut comptabilité"),
    ("remboursement", "rembourser validation client final montant remboursement"),
    ("rembourser", "remboursement procédure colis"),
    ("endommagé", "dommage colis incident photo constat déclaration"),
    ("fragile", "emballage responsabilité transport dommage"),
    ("adresse", "modification destination livraison changement"),
    ("destination", "adresse ville modification livraison"),
    ("téléphone", "contact téléphone client ramassage"),
    ("telephone", "téléphone coordonnées client contact modification livraison"),
    ("numéro", "téléphone contact client modification"),
    ("numero", "téléphone contact client modification"),
    ("supprimer", "annulation suppression colis annuler"),
    ("annuler", "annulation colis statut"),
    ("à préparer", "stock entrepôt préparation colis"),
    ("ramassage", "collecte pickup programmé demande ramassage"),
    ("ramassé", "colis statut annulation suppression annuler ramassage"),
    ("liste des villes", "couverture zone destination livrable"),
    ("plateforme", "assistance documentation tutoriel procédure"),
    ("aide", "assistance centre d'aide documentation SENDIT"),
    ("hors casablanca", "transport zone hub dommage responsabilité"),
    ("200 dh", "seuil validation remboursement montant dirhams"),
    ("livraison", "colis en cours tournée livreur statut modification coordonnées"),
    ("modifier", "modification changement coordonnées client procédure"),
    ("changement", "modification coordonnées téléphone adresse client"),
    ("coordonnées", "contact téléphone adresse client modification livraison"),
    ("coordonnees", "contact téléphone adresse client modification livraison"),
    ("déjà", "colis déjà statut livraison en cours tournée livreur"),
    ("deja", "colis statut livraison en cours tournée"),
    ("email", "courriel e-mail adresse messagerie modification procédure vendeur document pièce fournir"),
    ("e-mail", "email courriel modification adresse vendeur document"),
    ("courriel", "email e-mail modification vendeur document pièce"),
    ("messagerie", "email courriel e-mail modification compte"),
)

# Darija / arabizi tokens that rarely overlap French SOP wording — add FR lemmas for BM25.
_DARIJA_RETRIEVAL_HINTS: tuple[tuple[str, str], ...] = (
    ("takhzin", "stockage entrepôt magasin stock gratuit payant tarif prix frais"),
    ("tak7zin", "stockage entrepôt magasin stock gratuit payant tarif"),
    ("t7zin", "stockage entrepôt magasin stock inventaire"),
    ("takhzen", "stock stockage entrepôt"),
    ("fabor", "gratuit offert payant tarif prix frais"),
    ("fabour", "gratuit payant tarif prix"),
    ("blach", "gratuit payant tarif prix"),
)


def expand_query_for_retrieval_fr_darija(query: str) -> str:
    """Broaden BM25 recall: French/Darija hints + English→FR lemmas when EN terms match.

    English stays **without** blind French stuffing: only `_LOGISTICS_EN_FR` adds
    French when matching English substrings are present (vendor, delivery, …).
    """
    low = (query or "").lower()
    extra: List[str] = []
    for en, fr in _LOGISTICS_EN_FR:
        if en in low:
            extra.append(fr)
    for trigger, fr in _LOGISTICS_FR_HINTS:
        if trigger in low:
            extra.append(fr)
    for trigger, fr in _DARIJA_RETRIEVAL_HINTS:
        if trigger in low:
            extra.append(fr)
    if not extra and _DARIJA_LEX_BOOST.search(query or ""):
        extra.append(
            "livraison colis client téléphone coordonnées modification vendeur procédure"
        )
    if not extra:
        return query
    return f"{query}\n{' '.join(extra)}"


def _best_window_for_query(
    text: str,
    query: str,
    max_chars: int,
    *,
    expand_fr_darija_hints: bool = False,
) -> str:
    """When *text* is longer than *max_chars*, pick a slice that overlaps query terms / step numbers.

    Avoids always taking the document head, which hides later steps (e.g. §4.1).
    """
    if not text:
        return ""
    suffix = "\n…(tronqué)"
    top_note = "…(début du document omis)\n"
    if max_chars <= 120:
        return text[:max_chars] + ("…" if len(text) > max_chars else "")

    # Budget for inner slice after optional top note + suffix
    inner0 = max(120, max_chars - len(suffix))
    if len(text) <= inner0:
        return text

    q_use = (
        expand_query_for_retrieval_fr_darija(query)
        if expand_fr_darija_hints
        else (query or "")
    )
    q_tokens = [t for t in _tokenize(q_use) if len(t) >= 2]
    text_lower = text.lower()
    positions: List[int] = []
    step_positions: set[int] = set()

    for tok in set(q_tokens):
        pos = 0
        tlen = len(tok)
        while pos < len(text_lower):
            idx = text_lower.find(tok, pos)
            if idx == -1:
                break
            positions.append(idx)
            pos = idx + max(1, tlen)

    for m in re.finditer(r"\b\d+(?:\.\d+)+\b", query or ""):
        needle = m.group(0)
        pos = 0
        while True:
            idx = text.find(needle, pos)
            if idx == -1:
                break
            positions.append(idx)
            step_positions.add(idx)
            pos = idx + max(1, len(needle))

    for m in re.finditer(
        r"(?m)^[ \t]*(?:#{1,3}\s*)?(?:\d+(?:\.\d+)*)\s*(?:[.)-]|\s{1,3})",
        text[:300000],
    ):
        positions.append(m.start())

    positions = sorted(set(positions))
    upper0 = max(0, len(text) - inner0)
    step = max(400, inner0 // 5)

    def window_score(start: int, inner: int) -> float:
        frag = text[start : start + inner].lower()
        sc = 0.0
        for tok in set(q_tokens):
            c = frag.count(tok)
            if not c:
                continue
            sc += c * (2.5 if len(tok) >= 5 else 1.8 if len(tok) >= 4 else 1.0)
        return sc

    best_start = 0
    best_score = -1.0
    if positions:
        for center in positions:
            if center in step_positions:
                margin = min(420, max(180, inner0 // 3))
                s = max(0, min(center - margin, upper0))
            else:
                s = max(0, min(max(0, center - inner0 // 3), upper0))
            sc = window_score(s, inner0)
            if sc > best_score:
                best_score = sc
                best_start = s

    for s in range(0, upper0 + 1, step):
        sc = window_score(s, inner0)
        if sc > best_score:
            best_score = sc
            best_start = s

    if best_score <= 0.0 and upper0 > 0:
        best_start = min(upper0 // 3, upper0)

    use_prefix = best_start > 0
    inner = max(
        120,
        max_chars - len(suffix) - (len(top_note) if use_prefix else 0),
    )
    upper = max(0, len(text) - inner)
    best_start = min(best_start, upper)

    chunk = text[best_start : best_start + inner]
    if use_prefix:
        nl = chunk.find("\n")
        if 0 <= nl < 72:
            chunk = chunk[nl + 1 :]
    truncated_end = best_start + inner < len(text)
    if len(chunk) > inner:
        chunk = chunk[:inner]
    if truncated_end:
        cut = chunk.rsplit("\n", 1)[0]
        if cut:
            chunk = cut
    prefix = top_note if use_prefix else ""
    return prefix + chunk + (suffix if truncated_end else "")


def _greedy_inject_document_blocks(
    entries: List[tuple[str, str]],
    *,
    query: str,
    max_chars: int,
    expand_fr_darija_hints: bool,
) -> List[str]:
    """Prefer **full** top-ranked bodies; allow **at most one** query-aligned partial, then stop.

    Avoids allocating a thin slice to every file in the corpus (unreadable “Post-it” context).
    After the partial, still includes further files **only** if they fit entirely.
    """
    out: List[str] = []
    budget = max_chars
    used_partial = False
    q = (query or "").strip()
    min_partial = max(400, int(settings.RAG_MIN_CHARS_FOR_PARTIAL))

    for header, body in entries:
        overhead = len(header) + 1
        if budget <= overhead + 80:
            break
        max_body = budget - overhead
        if len(body) <= max_body:
            block = header + body + "\n"
            out.append(block)
            budget -= len(block)
            continue
        if used_partial:
            break
        if max_body < min_partial:
            break
        excerpt = _best_window_for_query(
            body,
            q,
            max_body,
            expand_fr_darija_hints=expand_fr_darija_hints,
        )
        block = header + excerpt + "\n"
        out.append(block)
        budget -= len(block)
        used_partial = True

    return out


def _collapse_ws(text: str) -> str:
    return collapse_whitespace(text)


def condense_sop_plaintext(text: str) -> str:
    """Tighten spacing/newlines for prompt budget without semantic rewriting."""
    if not text:
        return ""
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    return _collapse_ws(t)


def _read_txt(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8")
        return clean_sop_markdown(text)
    except Exception as exc:
        logger.warning("Could not read %s: %s", path.name, exc)
        return ""


def _read_md(path: Path) -> str:
    return _read_txt(path)


def _read_docx(path: Path) -> str:
    text = convert_docx_to_markdown(path)
    if text:
        return text
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())
        raw = "\n".join(parts)
        return clean_sop_markdown(raw)
    except Exception:
        try:
            import xml.etree.ElementTree as ET
            import zipfile

            with zipfile.ZipFile(path) as zf:
                with zf.open("word/document.xml") as f:
                    tree = ET.parse(f)
            raw = "\n".join(t.text for t in tree.iter() if t.tag.endswith("}t") and t.text)
            return clean_sop_markdown(raw)
        except Exception as exc:
            logger.warning("Could not read %s: %s", path.name, exc)
            return ""


def _read_pdf(path: Path) -> str:
    """Extract plain text from PDF (for ``data/documents/<cat>/pdf/*.pdf``)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning(
            "pypdf is not installed; skipping PDF %s (pip install pypdf)",
            path.name,
        )
        return ""
    try:
        reader = PdfReader(str(path))
        parts: List[str] = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t)
        raw = "\n\n".join(parts)
        return clean_sop_markdown(raw) if raw.strip() else ""
    except Exception as exc:
        logger.warning("Could not read PDF %s: %s", path.name, exc)
        return ""


@dataclass
class Doc:
    name: str
    category: str
    text: str
    tokens: List[str]
    tf: Counter


@dataclass
class CategoryIndex:
    name: str
    docs: List[Doc]
    df: Counter
    avgdl: float


class DocStore:
    """Per-category BM25 index. Each subfolder of data/documents/ is a category."""

    def __init__(
        self,
        docs_dir: Path = DOCS_DIR,
        docs_md_dir: Path = DOCS_MD_DIR,
        docs_txt_dir: Path = DOCS_TXT_DIR,
    ):
        self.docs_dir = docs_dir
        self.docs_md_dir = docs_md_dir
        self.docs_txt_dir = docs_txt_dir
        self.indexes: Dict[str, CategoryIndex] = {}
        self._load()

    def _load(self) -> None:
        if not self.docs_dir.is_dir():
            logger.warning("Documents dir not found: %s", self.docs_dir)
            return

        for sub in sorted(self.docs_dir.iterdir()):
            if not sub.is_dir():
                continue
            cat_name = sub.name
            md_cat = self.docs_md_dir / cat_name
            txt_cat = self.docs_txt_dir / cat_name
            docs: List[Doc] = []
            df: Counter = Counter()
            total_len = 0
            source = ""

            def ingest(name: str, text: str) -> None:
                nonlocal total_len
                text = append_to_document_text(cat_name, name, text)
                if not text.strip():
                    return
                toks = _tokenize(text)
                if not toks:
                    return
                tf = Counter(toks)
                docs.append(
                    Doc(name=name, category=cat_name, text=text, tokens=toks, tf=tf)
                )
                for term in tf:
                    df[term] += 1
                total_len += len(toks)

            # Prefer exported MD/TXT when they actually index at least one document.
            # If ``data/documents_md/<cat>`` contains only empty/broken .md files, we must
            # fall through to ``data/documents/<cat>/`` (Word/PDF) — otherwise help_md etc.
            # vanish from RAG despite a full docx corpus (common on pods).
            md_files = sorted(md_cat.rglob("*.md")) if md_cat.is_dir() else []
            if md_files:
                source = "documents_md"
                for p in md_files:
                    try:
                        rel = p.relative_to(md_cat)
                        name = rel.with_suffix("").as_posix()
                        ingest(name, _read_md(p))
                    except Exception:
                        logger.exception("DocStore: skip md %s", p.as_posix())
            if not docs and md_files:
                logger.warning(
                    "DocStore: category %r had %d .md path(s) under documents_md but indexed 0 docs; "
                    "trying documents_txt then data/documents (Word/PDF).",
                    cat_name,
                    len(md_files),
                )

            if not docs:
                txt_files = sorted(txt_cat.rglob("*.txt")) if txt_cat.is_dir() else []
                if txt_files:
                    source = "documents_txt"
                    for p in txt_files:
                        try:
                            rel = p.relative_to(txt_cat)
                            name = rel.with_suffix("").as_posix()
                            ingest(name, _read_txt(p))
                        except Exception:
                            logger.exception("DocStore: skip txt %s", p.as_posix())
                if not docs and txt_files:
                    logger.warning(
                        "DocStore: category %r had %d .txt path(s) under documents_txt but indexed 0 docs; "
                        "trying data/documents (Word/PDF).",
                        cat_name,
                        len(txt_files),
                    )

            if not docs:
                pdf_dir = sub / "pdf"
                file_jobs: List[tuple[Path, Any]] = []
                for p in sorted(sub.glob("*.docx")):
                    file_jobs.append((p, _read_docx))
                if pdf_dir.is_dir():
                    for p in sorted(pdf_dir.glob("*.pdf")):
                        file_jobs.append((p, _read_pdf))
                file_jobs.sort(key=lambda x: x[0].as_posix().lower())
                if file_jobs:
                    has_pdf = any(j[0].suffix.lower() == ".pdf" for j in file_jobs)
                    if has_pdf and any(j[0].suffix.lower() == ".docx" for j in file_jobs):
                        source = "docx+pdf"
                    elif has_pdf:
                        source = "pdf"
                    else:
                        source = "docx"
                    for p, reader_fn in file_jobs:
                        try:
                            try:
                                rel_stem = p.relative_to(sub).with_suffix("")
                                doc_name = rel_stem.as_posix()
                            except ValueError:
                                doc_name = p.stem
                            text = reader_fn(p)
                            ingest(doc_name, text)
                        except Exception:
                            logger.exception("DocStore: skip %s", p.as_posix())
            if docs:
                self.indexes[cat_name] = CategoryIndex(
                    name=cat_name,
                    docs=docs,
                    df=df,
                    avgdl=total_len / len(docs),
                )
                logger.info(
                    "DocStore: category '%s' loaded %d docs (%s)",
                    cat_name,
                    len(docs),
                    source,
                )
                if source == "docx":
                    logger.warning(
                        "DocStore: category '%s' - prefer `data/documents_md/%s/*.md` "
                        "(run `python -m scripts.export_sop_to_md`) for Markdown RAG",
                        cat_name,
                        cat_name,
                    )
            else:
                logger.warning("DocStore: category '%s' has no readable docs", cat_name)

        if not self.indexes:
            logger.warning("DocStore: no categories loaded from %s", self.docs_dir)

    def reload(self) -> None:
        """Rebuild BM25 indexes from disk (used after git pull or document export)."""
        self.indexes.clear()
        self._load()

    def list_categories(self) -> List[Dict]:
        return [
            {"name": c.name, "doc_count": len(c.docs), "doc_names": [d.name for d in c.docs]}
            for c in self.indexes.values()
        ]

    def rag_categories_all(self) -> List[str]:
        """Every indexed folder that has at least one document (sorted). Used for full-corpus chat RAG."""
        return sorted(name for name, idx in self.indexes.items() if idx.docs)

    def rag_categories_for_primary(self, primary: str) -> List[str]:
        """Primary chat category plus extras from RAG_EXTRA_CATEGORIES that exist on disk."""
        p = (primary or "").strip()
        if not p or p not in self.indexes:
            return []
        out: List[str] = [p]
        raw = (settings.RAG_EXTRA_CATEGORIES or "").strip()
        if not raw:
            return out
        for part in raw.split(","):
            e = part.strip()
            if not e or e in out:
                continue
            if e in self.indexes:
                out.append(e)
        return out

    def resolve_rag_scope(self, requested: Optional[str]) -> List[str]:
        """Retrieval scope from client: empty / 'all' → every indexed category; else one category (aliases)."""
        raw = (requested or "").strip().lower()
        if not raw or raw in ("all", "__all__", "*", "tout", "both"):
            return self.rag_categories_all()
        norm = raw.replace("-", "_")
        if norm in ("help", "help_md", "helpmd", "aide", "articles", "helpcenter"):
            key = "help_md"
        elif norm in ("procedures", "procedure", "sop", "sops"):
            key = "procedures"
        else:
            key = raw
        if key in self.indexes:
            return [key]
        logger.warning("RAG scope %r missing on disk — using all categories", requested)
        return self.rag_categories_all()

    def category_corpus_chars_multi(self, categories: List[str]) -> int:
        return sum(self.category_corpus_chars(c) for c in categories if c in self.indexes)

    def use_full_category_inject_multi(self, categories: List[str]) -> bool:
        if not categories:
            return False
        total_chars = self.category_corpus_chars_multi(categories)
        total_files = sum(len(self.indexes[c].docs) for c in categories if c in self.indexes)
        if total_files == 0:
            return False
        return (
            total_chars <= settings.RAG_FULL_CATEGORY_MAX_CHARS
            and total_files <= settings.RAG_FULL_CATEGORY_MAX_FILES
        )

    def _bm25(self, q_tokens: List[str], doc: Doc, idx: CategoryIndex,
              k1: float = 1.5, b: float = 0.75) -> float:
        N = len(idx.docs)
        dl = len(doc.tokens) or 1
        score = 0.0
        for term in q_tokens:
            df_t = idx.df.get(term, 0)
            if df_t == 0:
                continue
            idf = math.log((N - df_t + 0.5) / (df_t + 0.5) + 1.0)
            f = doc.tf.get(term, 0)
            denom = f + k1 * (1 - b + b * dl / (idx.avgdl or 1))
            score += idf * (f * (k1 + 1)) / (denom or 1)
        return score

    def category_corpus_chars(self, category: str) -> int:
        idx = self.indexes.get(category)
        if not idx:
            return 0
        return sum(len(d.text) for d in idx.docs)

    def use_full_category_inject(self, category: str) -> bool:
        """True → concatenate category docs (small corpus); False → BM25 top-k first."""
        idx = self.indexes.get(category)
        if not idx or not idx.docs:
            return False
        corpus = sum(len(d.text) for d in idx.docs)
        n = len(idx.docs)
        return corpus <= settings.RAG_FULL_CATEGORY_MAX_CHARS and n <= settings.RAG_FULL_CATEGORY_MAX_FILES

    def get_document_by_stem(self, category: str, stem: str) -> Optional[str]:
        """Return full document text for *stem* (filename without extension) or None."""
        idx = self.indexes.get(category)
        if not idx:
            return None
        for d in idx.docs:
            if d.name == stem:
                return d.text
        return None

    def get_document_by_catalog_id(self, catalog_id: str, fallback_category: str) -> Optional[str]:
        """Resolve agentic catalog id: ``category/stem`` or plain ``stem``."""
        raw = (catalog_id or "").strip()
        if not raw:
            return None
        if "/" in raw:
            cat, stem = raw.split("/", 1)
            cat, stem = cat.strip(), stem.strip()
            if cat in self.indexes and stem:
                return self.get_document_by_stem(cat, stem)
            return None
        return self.get_document_by_stem(fallback_category, raw)

    def retrieve(
        self,
        query: str,
        category: Optional[str] = None,
        *,
        categories: Optional[List[str]] = None,
        k: int = 5,
        expand_fr_darija_hints: bool = False,
    ) -> List[Doc]:
        if not self.indexes:
            return []
        if expand_fr_darija_hints:
            query = expand_query_for_retrieval_fr_darija(query)
        if categories:
            targets = [self.indexes[c] for c in categories if c in self.indexes]
        elif category and category in self.indexes:
            targets = [self.indexes[category]]
        else:
            targets = list(self.indexes.values())

        q_tokens = _tokenize(query)
        if not q_tokens:
            fallback_nt: List[Doc] = []
            for idx in targets:
                for d in idx.docs:
                    fallback_nt.append(d)
                    if len(fallback_nt) >= k:
                        return fallback_nt[:k]
            return fallback_nt

        if categories and len(targets) >= 2:
            floor = max(3, min(10, max(4, k // max(1, len(targets)))))
            seen: set[tuple[str, str]] = set()
            selected: List[Doc] = []
            for idx in targets:
                ranked = self._rank_docs_in_index(
                    idx, query, expand_for_retrieval=False
                )
                for d in ranked[:floor]:
                    key = (d.category, d.name)
                    if key in seen:
                        continue
                    seen.add(key)
                    selected.append(d)
            if len(selected) >= k:
                return selected[:k]
            scored_mc: List[tuple[float, Doc]] = []
            for idx in targets:
                for d in idx.docs:
                    key = (d.category, d.name)
                    if key in seen:
                        continue
                    s = self._bm25(q_tokens, d, idx)
                    if s > 0:
                        scored_mc.append((s, d))
            scored_mc.sort(key=lambda x: x[0], reverse=True)
            for d in (d for _, d in scored_mc):
                if len(selected) >= k:
                    break
                key = (d.category, d.name)
                if key in seen:
                    continue
                seen.add(key)
                selected.append(d)
            if selected:
                return selected

        scored = []
        for idx in targets:
            for d in idx.docs:
                s = self._bm25(q_tokens, d, idx)
                if s > 0:
                    scored.append((s, d))
        scored.sort(key=lambda x: x[0], reverse=True)
        out = [d for _, d in scored[:k]]
        if out:
            return out
        # Query terms missed the lexicon (scripts, typos): still return *some* docs.
        fallback: List[Doc] = []
        for idx in targets:
            for d in idx.docs:
                fallback.append(d)
                if len(fallback) >= k:
                    return fallback[:k]
        return fallback

    def _rank_docs_in_index(
        self,
        idx: CategoryIndex,
        query: str,
        *,
        expand_for_retrieval: bool,
    ) -> List[Doc]:
        """Stable sort: BM25 score desc, then original file order."""
        q = expand_query_for_retrieval_fr_darija(query) if expand_for_retrieval else query
        q_tokens = _tokenize(q)
        if not q_tokens:
            return list(idx.docs)
        scored: List[tuple[float, int, Doc]] = []
        for i, d in enumerate(idx.docs):
            s = self._bm25(q_tokens, d, idx)
            scored.append((s, i, d))
        scored.sort(key=lambda x: (-x[0], x[1]))
        return [t[2] for t in scored]

    def build_context(
        self,
        query: str,
        category: Optional[str] = None,
        k: int = 5,
        max_chars: int = 14000,
        *,
        categories: Optional[List[str]] = None,
        expand_fr_darija_hints: bool = False,
        condense: bool = False,
    ) -> str:
        top = self.retrieve(
            query,
            category=category,
            categories=categories,
            k=k,
            expand_fr_darija_hints=expand_fr_darija_hints,
        )
        if not top:
            return ""
        entries: List[tuple[str, str]] = []
        for d in top:
            body = condense_sop_plaintext(d.text) if condense else d.text.strip()
            header = f"### Document : {d.name}  (catégorie : {d.category})\n"
            entries.append((header, body))
        if settings.RAG_GREEDY_FULL_DOCS:
            blocks = _greedy_inject_document_blocks(
                entries,
                query=query,
                max_chars=max_chars,
                expand_fr_darija_hints=expand_fr_darija_hints,
            )
            return "\n".join(blocks)
        parts: List[str] = []
        budget = max_chars
        for header, body in entries:
            overhead = len(header) + 1
            if budget <= overhead + 80:
                break
            max_body = budget - overhead
            if len(body) <= max_body:
                block = header + body + "\n"
                parts.append(block)
                budget -= len(block)
            else:
                excerpt = _best_window_for_query(
                    body,
                    query,
                    max_body,
                    expand_fr_darija_hints=expand_fr_darija_hints,
                )
                block = header + excerpt + "\n"
                parts.append(block)
                budget -= len(block)
            if budget <= 0:
                break
        return "\n".join(parts)

    def build_all_docs_context(
        self,
        category: Optional[str] = None,
        max_chars: int = 15000,
        *,
        categories: Optional[List[str]] = None,
        query: Optional[str] = None,
        expand_for_retrieval: bool = False,
        condense: bool = False,
    ) -> str:
        """Concatenate documents in one or more categories up to *max_chars*."""
        if not self.indexes:
            return ""
        if categories:
            targets = [self.indexes[c] for c in categories if c in self.indexes]
        elif category and category in self.indexes:
            targets = [self.indexes[category]]
        else:
            targets = list(self.indexes.values())

        parts: List[str] = []
        used = 0
        remaining = max_chars
        for idx in targets:
            if remaining <= 100:
                break
            docs = (
                self._rank_docs_in_index(idx, query, expand_for_retrieval=expand_for_retrieval)
                if (query and query.strip())
                else list(idx.docs)
            )
            if not docs:
                continue
            n = len(docs)
            entries: List[tuple[str, str, str]] = []
            for d in docs:
                body = condense_sop_plaintext(d.text) if condense else d.text.strip()
                header = f"### Document : {d.name}  (catégorie : {d.category})\n"
                entries.append((header, body, d.name))

            # Separator between doc blocks: single newline after body.
            # *remaining* is shared across all categories (multi-corpus inject must not
            # re-use the full max_chars budget per folder — that blew past vLLM context).
            overhead = sum(len(h) + 1 for h, _, _ in entries)
            budget_bodies = max(0, remaining - overhead - max(0, n - 1))
            total_body = sum(len(b) for _, b, _ in entries)

            if total_body == 0:
                for header, _, _name in entries:
                    block = header + "\n"
                    if len(block) > remaining:
                        break
                    parts.append(block)
                    used += len(block)
                    remaining -= len(block)
                continue

            if total_body <= budget_bodies:
                for header, body, _name in entries:
                    block = header + body + "\n"
                    if len(block) > remaining:
                        block = block[:remaining].rsplit("\n", 1)[0] + "\n…(tronqué)\n"
                    parts.append(block)
                    u = len(block)
                    used += u
                    remaining -= u
                    if remaining <= 100:
                        break
                continue

            if settings.RAG_GREEDY_FULL_DOCS:
                blocks = _greedy_inject_document_blocks(
                    [(h, b) for h, b, _ in entries],
                    query=query or "",
                    max_chars=remaining,
                    expand_fr_darija_hints=expand_for_retrieval,
                )
                for b in blocks:
                    if remaining <= 0:
                        break
                    if len(b) > remaining:
                        b = b[: max(0, remaining)].rsplit("\n", 1)[0] + "\n…(tronqué)\n"
                    parts.append(b)
                    u = len(b)
                    used += u
                    remaining -= u
                continue

            if query and query.strip():
                # Relevance order already applied in *docs*; give top hits most budget so
                # one procedure is readable instead of slicing every file equally thin.
                exp = 1.25
                raw_w = [1.0 / ((i + 1) ** exp) for i in range(n)]
                sw = sum(raw_w)
                alloc = [int(budget_bodies * (w / sw)) for w in raw_w]
                drift = budget_bodies - sum(alloc)
                if drift != 0 and alloc:
                    alloc[0] += drift
            else:
                ratios = [len(b) / total_body for _, b, _ in entries]
                raw_alloc = [int(budget_bodies * r) for r in ratios]
                alloc = raw_alloc
                diff = budget_bodies - sum(alloc)
                if diff != 0 and alloc:
                    alloc[-1] = max(0, alloc[-1] + diff)

            for i, (header, body, _name) in enumerate(entries):
                if remaining <= 100:
                    break
                cap = alloc[i] if i < len(alloc) else 0
                if len(body) <= cap:
                    chunk = body
                    suffix = "\n"
                else:
                    if query and query.strip() and cap > 120:
                        chunk = _best_window_for_query(
                            body,
                            query,
                            cap,
                            expand_fr_darija_hints=expand_for_retrieval,
                        )
                        suffix = "\n"
                    else:
                        chunk = body[:cap].rsplit("\n", 1)[0] if cap > 80 else body[:cap]
                        chunk = chunk or body[:cap]
                        suffix = "\n…(tronqué — budget contexte atteint)\n"
                block = header + chunk + suffix
                if len(block) > remaining:
                    over = len(block) - remaining
                    if over > 0 and len(chunk) > over + 20:
                        chunk = chunk[: max(0, len(chunk) - over - 20)].rsplit("\n", 1)[0]
                        block = header + chunk + "\n…(tronqué)\n"
                    else:
                        block = block[:remaining]
                parts.append(block)
                u = len(block)
                used += u
                remaining -= u

        logger.info(
            "Injected %d document blocks into context (~%d chars of %d budget, condense=%s)",
            len(parts),
            used,
            max_chars,
            condense,
        )
        return "\n".join(parts)


_store: Optional[DocStore] = None


def get_store() -> DocStore:
    global _store
    if _store is None:
        _store = DocStore()
    return _store


def reload_document_store() -> None:
    """Reload RAG indices from disk without restarting the API process."""
    global _store
    if _store is None:
        _store = DocStore()
    else:
        _store.reload()

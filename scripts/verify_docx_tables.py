#!/usr/bin/env python3
"""Verify how Word tables are interpreted for RAG (.docx → Markdown).

Without arguments: runs a synthetic .docx round-trip (requires python-docx).

With paths: prints full extracted text for each file (inspect markdown pipes for tables).

Uses ``core.docx_to_md.convert_docx_to_markdown``. Limitations:
- First table row is always a Markdown *header* row.
- Heavily merged cells may misalign vs Word’s grid.
- Metadata-style tables (titre, référence, …) are omitted by design.
"""
from __future__ import annotations

import argparse
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from core.docx_to_md import convert_docx_to_markdown  # noqa: E402


def _synthetic_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Paragraph before table.")
    t = doc.add_table(rows=3, cols=3)
    for j, h in enumerate(["Étape", "Action", "Délai"]):
        t.cell(0, j).text = h
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "Contacter le client"
    t.cell(1, 2).text = "24h"
    t.cell(2, 0).text = "2"
    t.cell(2, 1).text = "Mettre à jour le système"
    t.cell(2, 2).text = "48h"
    doc.add_paragraph("Paragraph after table.")
    doc.save(path)


def run_self_test() -> None:
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "synthetic_sop.docx"
        _synthetic_docx(p)
        text = convert_docx_to_markdown(p)
    if "Paragraph before" not in text or "Paragraph after" not in text:
        raise SystemExit("FAIL: paragraphs missing from ordered parse")
    if "Contacter le client" not in text or "Mettre à jour" not in text:
        raise SystemExit("FAIL: table cell text missing")
    pipe_lines = [ln for ln in text.splitlines() if ln.strip().startswith("|")]
    if len(pipe_lines) < 3:
        raise SystemExit(f"FAIL: expected markdown table lines, got: {pipe_lines!r}")
    print("Self-test OK: synthetic docx paragraphs + table cells round-trip correctly.")


def _table_blocks(text: str) -> list[str]:
    blocks: list[list[str]] = []
    cur: list[str] = []
    for line in text.splitlines():
        if line.strip().startswith("|"):
            cur.append(line)
        else:
            if cur:
                blocks.append(cur)
                cur = []
    if cur:
        blocks.append(cur)
    return ["\n".join(b) for b in blocks]


def print_report(path: Path, show_raw: bool) -> None:
    text = convert_docx_to_markdown(path)
    if not text.strip():
        print(f"{path}: (empty or unreadable)")
        return
    print(f"\n{'=' * 72}\nFILE: {path}\n{'=' * 72}")
    n_tables = len(_table_blocks(text))
    print(f"Approx. markdown table block(s): {n_tables}  |  total chars: {len(text)}")
    if show_raw:
        print("--- extracted text ---")
        print(text)
        print("--- end ---")
    else:
        for i, block in enumerate(_table_blocks(text), 1):
            print(f"\n--- table block {i} ({block.count(chr(10)) + 1} lines) ---")
            print(block)
        non_table = "\n".join(
            ln for ln in text.splitlines() if not ln.strip().startswith("|")
        )
        preview = non_table.strip()
        if preview:
            print("\n--- non-table preview (first 1200 chars) ---")
            print(preview[:1200] + ("…" if len(preview) > 1200 else ""))


def main() -> None:
    ap = argparse.ArgumentParser(description="Verify .docx table extraction for RAG")
    ap.add_argument(
        "docx",
        nargs="*",
        type=Path,
        help=".docx file(s); if omitted, run synthetic self-test only",
    )
    ap.add_argument(
        "--raw",
        action="store_true",
        help="Print full extracted text (including paragraphs)",
    )
    args = ap.parse_args()

    if not args.docx:
        run_self_test()
        return

    missing = [p for p in args.docx if not p.is_file()]
    if missing:
        for p in missing:
            print(f"Not found: {p}", file=sys.stderr)
        raise SystemExit(1)

    for p in args.docx:
        print_report(p.resolve(), show_raw=args.raw)


if __name__ == "__main__":
    main()

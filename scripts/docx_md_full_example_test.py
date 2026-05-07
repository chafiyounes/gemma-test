#!/usr/bin/env python3
"""Build a rich synthetic .docx and assert Markdown conversion keeps full body context.

Run from repo root:
    python scripts/docx_md_full_example_test.py
Requires: python-docx
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from core.docx_to_md import convert_docx_to_markdown  # noqa: E402


def main() -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise SystemExit(f"python-docx required: {e}") from e

    doc = Document()
    doc.add_heading("SOP exemple complet", level=1)
    doc.add_paragraph(
        "Ce paragraphe suit le titre et doit apparaître dans le Markdown, "
        "avec le corps du document (pas seulement les tableaux)."
    )
    doc.add_paragraph("Étape intro avant liste.", style=doc.styles["Normal"])

    try:
        bullet_style = doc.styles["List Bullet"]
        doc.add_paragraph("Point liste A", style=bullet_style)
        doc.add_paragraph("Point liste B", style=bullet_style)
    except KeyError:
        p = doc.add_paragraph()
        p.style = "List Paragraph"
        p.add_run("Point liste (fallback style)")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Colonne A"
    table.cell(0, 1).text = "Colonne B"
    table.cell(1, 0).text = "valeur 1"
    table.cell(1, 1).text = "valeur 2"
    table.cell(2, 0).text = "ligne avec | pipe"
    table.cell(2, 1).text = "échappé attendu"

    outer = doc.add_table(rows=1, cols=1)
    oc = outer.cell(0, 0)
    oc.text = "Cellule avec table imbriquée"
    inner = oc.add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "in-a"
    inner.cell(0, 1).text = "in-b"
    inner.cell(1, 0).text = "in-c"
    inner.cell(1, 1).text = "in-d"

    doc.add_paragraph("Paragraphe final après tous les blocs.")

    with tempfile.TemporaryDirectory() as td:
        path = Path(td) / "full_example.docx"
        doc.save(path)
        md = convert_docx_to_markdown(path)

    must = [
        "SOP exemple complet",
        "Ce paragraphe suit le titre",
        "Colonne A",
        "Colonne B",
        "valeur 1",
        "pipe",
        "Cellule avec table imbriquée",
        "in-a",
        "in-d",
        "Paragraphe final",
    ]
    missing = [s for s in must if s.lower() not in md.lower()]
    if missing:
        print("FAIL — fragments manquants:", missing)
        print("--- markdown ---\n", md[:4000])
        raise SystemExit(1)

    if "|" not in md or md.count("|") < 6:
        print("FAIL — pas assez de syntaxe tableau Markdown\n", md)
        raise SystemExit(1)

    if "#" not in md:
        print("FAIL — titre Markdown attendu\n", md)
        raise SystemExit(1)

    print("OK — document Word complet converti (titres, paragraphes, table, table imbriquée).")
    print(f"({len(md)} caractères générés)")


if __name__ == "__main__":
    main()
